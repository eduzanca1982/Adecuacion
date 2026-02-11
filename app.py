# app.py — Nano Opal v25.0
# Objetivo: generación por alumno (Sheets) + DOCX alumno/docente + ZIP.
# Cambios clave:
# - Modelo de texto FIJO: gemini-2.5-flash (sin selector lateral)
# - Sin “on-the-fly” ni controles manuales de modelos
# - Prompt por alumno SIEMPRE incluye Grupo + Diagnóstico + Grado
# - JSON: tolerante (autofill, normalización, repair/retry, fallback fuerte anti-archivos vacíos)
# - Crear: textarea + botón “Enviar prompt” (st.form), no Ctrl+Enter
# - ZIP incluye _REPORTE.txt + _RESUMEN.txt

import streamlit as st
import google.generativeai as genai
import pandas as pd
import json
import io
import zipfile
import time
import random
import hashlib
import base64
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# Config
# ============================================================
st.set_page_config(page_title="Nano Opal v25.0 🍌", layout="wide", page_icon="🍌")

SHEET_ID = "1dCZdGmK765ceVwTqXzEAJCrdSvdNLBw7t3q5Cq1Qrww"
URL_PLANILLA = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/export?format=csv"

DEFAULT_TEXT_MODEL = "models/gemini-2.5-flash"  # FIJO
IMG_MODEL_CANDIDATES = [
    "models/gemini-2.5-flash-image",
    "models/gemini-2.0-flash-image",
    "models/imagen-3.0-generate-002",
    "models/imagen-3.0-generate-001",
    "models/imagen-3.0-generate-000",
]

RETRIES = 6
CACHE_TTL_SECONDS = 6 * 60 * 60

MIN_IMAGE_BYTES = 1200
IMAGE_PROMPT_PREFIX = "Pictograma estilo ARASAAC, trazos negros gruesos, fondo blanco, ultra simple, sin sombras de: "

SAFETY_SETTINGS = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
]

BASE_GEN_CFG_JSON = {
    "response_mime_type": "application/json",
    "temperature": 0,
    "top_p": 1,
    "top_k": 1,
}

OUT_TOKEN_STEPS = [4096, 6144, 8192]
OUT_TOKEN_STEPS_FALLBACK = [2048, 3072, 4096]

ACTION_EMOJI_BY_TIPO = {
    "completar": "✍️",
    "escritura": "✍️",
    "multiple choice": "🔢",
    "multiple_choice": "🔢",
    "seleccion": "🔢",
    "unir": "📖",
    "lectura": "📖",
    "verdadero_falso": "📖",
    "problema_guiado": "🔢",
    "calcular": "🔢",
    "dibujar": "🎨",
    "arte": "🎨",
}

SYSTEM_PROMPT = f"""
Actúa como un Senior Inclusive UX Designer y Tutor Psicopedagogo.

Objetivo: producir una ficha de 60 minutos neuroinclusiva (TDAH/dislexia friendly) con estética tipo "Card"
y producir un solucionario para el docente.

REGLAS NO NEGOCIABLES:
- SALIDA: JSON puro, sin texto extra.
- NO uses markdown. NO uses ** ni __ ni backticks. CERO marcadores de negrita.
- ICONOS: Cada ítem en items[] debe iniciar el enunciado con un emoji de acción:
  ✍️ completar/escribir, 📖 leer, 🔢 calcular, 🎨 dibujar.
- MICRO-PASOS: pista_visual debe ser andamiaje físico/visual, instrucciones concretas. No teoría.
- LENGUAJE: 1 acción por frase, pasos numerados cuando aplique.
- VISUAL: si visual.habilitado=true, visual.prompt debe comenzar EXACTAMENTE con:
  "{IMAGE_PROMPT_PREFIX}[OBJETO]"

ESQUEMA (no lo rompas):
{{
  "objetivo_aprendizaje": "string",
  "tiempo_total_min": 60,
  "consigna_general_alumno": "string (paso a paso, sin saludos)",
  "items": [
    {{
      "tipo": "calcular|lectura|escritura|dibujar|multiple choice|unir|completar|verdadero_falso|problema_guiado",
      "enunciado": "string (DEBE EMPEZAR con emoji de acción)",
      "pasos": ["string","string"],
      "opciones": ["string","string"],
      "respuesta_formato": "texto_corto|procedimiento|dibujo|multiple_choice",
      "keywords_bold": ["string","string"],
      "pista_visual": "string (micro-pasos concretos)",
      "visual": {{ "habilitado": boolean, "prompt": "string" }}
    }}
  ],
  "adecuaciones_aplicadas": ["string","string"],
  "sugerencias_docente": ["string","string"],
  "solucionario_docente": {{
    "respuestas": [
      {{
        "item_index": 1,
        "respuesta_final": "string",
        "desarrollo": ["string","string"],
        "errores_frecuentes": ["string","string"]
      }}
    ],
    "criterios_correccion": ["string","string"]
  }},
  "control_calidad": {{
    "items_count": number,
    "incluye_ejemplo": boolean,
    "lenguaje_concreto": boolean,
    "una_accion_por_frase": boolean,
    "sin_markdown": boolean
  }}
}}
""".strip()


# ============================================================
# Helpers
# ============================================================
def now_str() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def hash_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()


def safe_filename(name: str) -> str:
    s = str(name).strip().replace(" ", "_")
    for ch in ["/", "\\", ":", "*", "?", "\"", "<", ">", "|"]:
        s = s.replace(ch, "_")
    while "__" in s:
        s = s.replace("__", "_")
    return (s or "SIN_NOMBRE")[:120]


def _is_retryable_error(e: Exception) -> bool:
    s = str(e).lower()
    markers = [
        "429", "too many requests", "rate", "quota", "resource exhausted",
        "timeout", "timed out", "deadline", "unavailable", "503", "500", "internal",
        "connection reset", "temporarily", "service unavailable",
    ]
    return any(m in s for m in markers)


def retry_with_backoff(fn):
    last = None
    for i in range(RETRIES):
        try:
            return fn()
        except Exception as e:
            last = e
            if i == RETRIES - 1 or not _is_retryable_error(e):
                raise
            sleep = (2 ** i) + random.uniform(0, 0.75)
            time.sleep(min(sleep, 25))
    raise last


def normalize_bool(v: Any) -> bool:
    if isinstance(v, bool):
        return v
    if isinstance(v, (int, float)):
        return bool(v)
    if isinstance(v, str):
        return v.strip().lower() in {"true", "1", "yes", "y", "si", "sí"}
    return False


def _contains_markdown_markers(s: str) -> bool:
    if not s:
        return False
    return ("**" in s) or ("```" in s) or ("__" in s)


def ensure_action_emoji(tipo: str, enunciado: str) -> str:
    t = (tipo or "").strip().lower()
    e = (enunciado or "").strip()
    if not e:
        return e
    if any(e.startswith(x) for x in ["✍️", "📖", "🔢", "🎨"]):
        return e
    emoji = ACTION_EMOJI_BY_TIPO.get(t, "📖")
    return f"{emoji} {e}"


def normalize_visual_prompt(p: str) -> str:
    p = (p or "").strip()
    if not p:
        return p
    if p.startswith(IMAGE_PROMPT_PREFIX):
        return p
    return IMAGE_PROMPT_PREFIX + p


def validate_text_input(text: str, mode: str) -> Tuple[bool, str, Dict[str, Any]]:
    info = {
        "chars": len(text or ""),
        "lines": (text or "").count("\n") + (1 if text else 0),
        "preview": (text or "")[:1600],
    }
    if mode == "ADAPTAR":
        if not text or not text.strip():
            return False, "TEXTO vacío tras extracción.", info
        if len(text) < 120:
            return False, "TEXTO muy corto (<120 chars).", info
        return True, "OK", info
    if not text or not text.strip():
        return False, "Prompt vacío.", info
    return True, "OK", info


# ============================================================
# JSON sanitization + parsing
# ============================================================
_JSON_CODEFENCE_RE = re.compile(r"```(?:json)?\s*([\s\S]*?)\s*```", re.IGNORECASE)
_TRAILING_COMMA_OBJ_RE = re.compile(r",\s*}")
_TRAILING_COMMA_ARR_RE = re.compile(r",\s*]")
_SINGLELINE_COMMENT_RE = re.compile(r"//.*?$", re.MULTILINE)
_BLOCK_COMMENT_RE = re.compile(r"/\*[\s\S]*?\*/", re.MULTILINE)
_UNQUOTED_KEY_RE = re.compile(r'(\{|,)\s*([A-Za-z_][A-Za-z0-9_ ]{0,60}?)\s*:')


def _strip_wrappers(raw: str) -> str:
    if not raw:
        return raw
    m = _JSON_CODEFENCE_RE.search(raw)
    if m:
        raw = m.group(1)
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return raw.strip()
    return raw[start:end + 1].strip()


def _fix_unquoted_keys(s: str) -> str:
    def repl(m):
        prefix = m.group(1)
        key = m.group(2).strip()
        key = re.sub(r"\s+", "_", key)
        return f'{prefix} "{key}":'
    return _UNQUOTED_KEY_RE.sub(repl, s)


def sanitize_json_text(raw_text: str) -> str:
    if not raw_text:
        return raw_text
    s = raw_text
    s = _strip_wrappers(s)
    s = _BLOCK_COMMENT_RE.sub("", s)
    s = _SINGLELINE_COMMENT_RE.sub("", s)
    s = _TRAILING_COMMA_OBJ_RE.sub("}", s)
    s = _TRAILING_COMMA_ARR_RE.sub("]", s)
    s = _fix_unquoted_keys(s)
    return s.strip()


def safe_json_loads(raw_text: str) -> Dict[str, Any]:
    if not raw_text:
        raise ValueError("Empty JSON text")
    cleaned = sanitize_json_text(raw_text)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        cleaned2 = _strip_wrappers(cleaned)
        cleaned2 = _TRAILING_COMMA_OBJ_RE.sub("}", cleaned2)
        cleaned2 = _TRAILING_COMMA_ARR_RE.sub("]", cleaned2)
        cleaned2 = _fix_unquoted_keys(cleaned2)
        try:
            return json.loads(cleaned2)
        except json.JSONDecodeError as e2:
            raise ValueError(f"JSON parse failed: {e} // {e2}")


def build_parse_fix_prompt(raw: str, err: str) -> str:
    return f"""
Tu única tarea es convertir el siguiente texto en JSON válido.
No agregues texto. No agregues comentarios.
No cambies el contenido semántico, solo corrige sintaxis JSON.
Reglas: keys con comillas dobles, sin trailing commas, valores string con comillas dobles.

ERROR DE PARSEO:
{err}

TEXTO A CONVERTIR:
{raw}
""".strip()


def build_repair_prompt(raw: str, why: str) -> str:
    return f"""
Devuelve EXCLUSIVAMENTE un JSON válido del esquema (sin texto extra).
No cambies el contenido pedagógico salvo lo necesario para cumplir el esquema y reglas.

Problema detectado:
{why}

JSON A CORREGIR:
{raw}

Reglas:
- Prohibido markdown. NO usar ** ni backticks ni __.
- TODAS las keys deben ir entre comillas dobles.
- Sin trailing commas.
- tiempo_total_min = 60
- items[] debe tener 6 items.
- items[].enunciado inicia con emoji (✍️📖🔢🎨)
- visual.prompt inicia con "{IMAGE_PROMPT_PREFIX}" si visual.habilitado=true
""".strip()


# ============================================================
# DOCX extraction (párrafos + tablas)
# ============================================================
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _extract_text_from_el(el) -> str:
    return "".join(n.text for n in el.iter() if n.tag == f"{W_NS}t" and n.text).strip()


def extraer_texto_docx(file) -> str:
    doc = Document(file)
    out: List[str] = []
    for el in doc.element.body:
        if el.tag == f"{W_NS}p":
            t = _extract_text_from_el(el)
            if t:
                out.append(t)
        elif el.tag == f"{W_NS}tbl":
            # nota: recorrer filas/td dentro del tbl actual
            for row in el.findall(f".//{W_NS}tr"):
                cells = [_extract_text_from_el(c) for c in row.findall(f".//{W_NS}tc")]
                if any(cells):
                    out.append(" | ".join(cells))
            out.append("")
    return "\n".join(out).strip()


# ============================================================
# Gemini response parsing
# ============================================================
def _extract_text_or_none(resp) -> Optional[str]:
    try:
        cand = resp.candidates[0]
        content = getattr(cand, "content", None)
        if not content or not getattr(content, "parts", None):
            return None
        chunks = []
        for p in content.parts:
            t = getattr(p, "text", None)
            if t:
                chunks.append(t)
        out = "".join(chunks).strip()
        return out if out else None
    except Exception:
        return None


def _finish_reason(resp) -> Optional[int]:
    try:
        return int(resp.candidates[0].finish_reason)
    except Exception:
        return None


# ============================================================
# Image parsing (best-effort)
# ============================================================
DATA_URI_RE = re.compile(r"data:image/(png|jpeg|jpg|webp);base64,([A-Za-z0-9+/=\n\r]+)")


def _maybe_b64_to_bytes(x: Any) -> Optional[bytes]:
    if x is None:
        return None
    if isinstance(x, (bytes, bytearray)):
        return bytes(x)
    if isinstance(x, str):
        s = x.strip()
        m = DATA_URI_RE.search(s)
        if m:
            b64 = m.group(2)
            try:
                return base64.b64decode(b64, validate=False)
            except Exception:
                return None
        if len(s) > 400 and all(c in "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/=\n\r" for c in s[:200]):
            try:
                return base64.b64decode(s, validate=False)
            except Exception:
                return None
    return None


def _extract_inline_bytes_or_none(resp) -> Optional[bytes]:
    try:
        cand = resp.candidates[0]
        content = getattr(cand, "content", None)
        if not content or not getattr(content, "parts", None):
            return None
        for part in content.parts:
            inline = getattr(part, "inline_data", None)
            if inline is not None:
                data = getattr(inline, "data", None)
                b = _maybe_b64_to_bytes(data) or (data if isinstance(data, (bytes, bytearray)) else None)
                if b:
                    return b
            inline2 = getattr(part, "inlineData", None)
            if inline2 is not None:
                data2 = getattr(inline2, "data", None)
                b2 = _maybe_b64_to_bytes(data2)
                if b2:
                    return b2
            t = getattr(part, "text", None)
            b3 = _maybe_b64_to_bytes(t)
            if b3:
                return b3
        return None
    except Exception:
        return None


def _looks_like_image(b: bytes) -> bool:
    if not b or len(b) < MIN_IMAGE_BYTES:
        return False
    if b[:8] == b"\x89PNG\r\n\x1a\n":
        return True
    if b[:3] == b"\xff\xd8\xff":
        return True
    if b[:4] == b"RIFF" and b[8:12] == b"WEBP":
        return True
    return False


def generate_image_bytes(model_id: str, prompt_img: str) -> Optional[bytes]:
    if not model_id:
        return None
    prompt_img = normalize_visual_prompt(prompt_img)

    def call_with_cfg(cfg: Optional[Dict[str, Any]]):
        m = genai.GenerativeModel(model_id)
        if cfg is None:
            return m.generate_content(prompt_img, safety_settings=SAFETY_SETTINGS)
        return m.generate_content(prompt_img, generation_config=cfg, safety_settings=SAFETY_SETTINGS)

    cfg_variants = [
        {"response_modalities": ["Image"]},
        {"response_modalities": ["IMAGE"]},
        {"responseModalities": ["Image"]},
        {"responseModalities": ["IMAGE"]},
        None,
    ]

    for cfg in cfg_variants:
        try:
            resp = retry_with_backoff(lambda: call_with_cfg(cfg))
            b = _extract_inline_bytes_or_none(resp)
            if b and _looks_like_image(b):
                return b
        except Exception:
            continue
    return None


@st.cache_resource(show_spinner=False)
def pick_image_model() -> Optional[str]:
    # no UI; autodetect best-effort
    for mid in IMG_MODEL_CANDIDATES:
        try:
            b = generate_image_bytes(mid, IMAGE_PROMPT_PREFIX + "manzana")
            if b and _looks_like_image(b):
                return mid
        except Exception:
            continue
    return None


def smoke_test_text_model(model_id: str) -> Tuple[bool, str]:
    try:
        m = genai.GenerativeModel(model_id)
        cfg = {"temperature": 0, "max_output_tokens": 32}
        resp = retry_with_backoff(lambda: m.generate_content("Responde SOLO: OK", generation_config=cfg, safety_settings=SAFETY_SETTINGS))
        t = _extract_text_or_none(resp)
        if not t:
            fr = _finish_reason(resp)
            return False, f"Sin texto (finish_reason={fr})"
        return True, "OK"
    except Exception as e:
        return False, f"{type(e).__name__}: {e}"


# ============================================================
# Normalización “anti-rigidez” (autofill) + anti-archivo vacío
# ============================================================
def _clean_str(x: Any) -> str:
    s = "" if x is None else str(x)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    return s.strip()


def _coerce_list_str(x: Any, max_n: int) -> List[str]:
    if isinstance(x, list):
        out = []
        for v in x:
            if v is None:
                continue
            sv = str(v).strip()
            if sv:
                out.append(sv)
        return out[:max_n]
    if isinstance(x, str) and x.strip():
        return [x.strip()][:max_n]
    return []


def normalize_activity_payload(data: Dict[str, Any]) -> Dict[str, Any]:
    # Tolerante: crea claves faltantes y arregla formatos.
    if not isinstance(data, dict):
        data = {}

    # root defaults
    data.setdefault("objetivo_aprendizaje", "")
    data.setdefault("tiempo_total_min", 60)
    data.setdefault("consigna_general_alumno", "")
    data.setdefault("items", [])
    data.setdefault("adecuaciones_aplicadas", [])
    data.setdefault("sugerencias_docente", [])
    data.setdefault("solucionario_docente", {})
    data.setdefault("control_calidad", {})

    data["objetivo_aprendizaje"] = _clean_str(data.get("objetivo_aprendizaje"))
    data["consigna_general_alumno"] = _clean_str(data.get("consigna_general_alumno"))
    data["tiempo_total_min"] = 60

    # lists
    data["adecuaciones_aplicadas"] = _coerce_list_str(data.get("adecuaciones_aplicadas"), 30)
    data["sugerencias_docente"] = _coerce_list_str(data.get("sugerencias_docente"), 30)

    # items normalize
    items_in = data.get("items", [])
    if not isinstance(items_in, list):
        items_in = []

    items_norm: List[Dict[str, Any]] = []
    for it in items_in[:50]:
        if not isinstance(it, dict):
            continue

        tipo = _clean_str(it.get("tipo")) or "lectura"
        en = ensure_action_emoji(tipo, _clean_str(it.get("enunciado")))
        pasos = _coerce_list_str(it.get("pasos"), 10)
        opciones = _coerce_list_str(it.get("opciones"), 10)
        resp_fmt = _clean_str(it.get("respuesta_formato")) or "texto_corto"
        kw = _coerce_list_str(it.get("keywords_bold"), 12)
        pista = _clean_str(it.get("pista_visual"))

        v = it.get("visual", {})
        if not isinstance(v, dict):
            v = {}
        v_en = normalize_bool(v.get("habilitado", False))
        v_pr = normalize_visual_prompt(_clean_str(v.get("prompt"))) if v_en else ""

        # hard guard: no markdown markers
        if _contains_markdown_markers(en):
            en = en.replace("**", "").replace("__", "").replace("```", "")
        if _contains_markdown_markers(pista):
            pista = pista.replace("**", "").replace("__", "").replace("```", "")

        items_norm.append({
            "tipo": tipo,
            "enunciado": en,
            "pasos": pasos,
            "opciones": opciones,
            "respuesta_formato": resp_fmt,
            "keywords_bold": kw,
            "pista_visual": pista,
            "visual": {"habilitado": bool(v_en), "prompt": v_pr},
        })

    # enforce 6 items minimum (anti-archivo vacío)
    while len(items_norm) < 6:
        k = len(items_norm) + 1
        items_norm.append({
            "tipo": "lectura",
            "enunciado": f"📖 Lee y responde (Ítem {k}).",
            "pasos": ["Lee despacio.", "Responde con una palabra."],
            "opciones": [],
            "respuesta_formato": "texto_corto",
            "keywords_bold": ["lee", "responde"],
            "pista_visual": "Haz un paso por vez.",
            "visual": {"habilitado": True, "prompt": IMAGE_PROMPT_PREFIX + "libro"},
        })

    data["items"] = items_norm[:12]

    # solucionario normalize
    sol = data.get("solucionario_docente", {})
    if not isinstance(sol, dict):
        sol = {}
    resp = sol.get("respuestas", [])
    if not isinstance(resp, list):
        resp = []

    by_idx: Dict[int, Dict[str, Any]] = {}
    for r in resp:
        if not isinstance(r, dict):
            continue
        try:
            idx = int(r.get("item_index", 0) or 0)
        except Exception:
            idx = 0
        if idx <= 0:
            continue
        by_idx[idx] = r

    respuestas_out = []
    for i in range(1, len(data["items"]) + 1):
        r = by_idx.get(i, {})
        rf = _clean_str(r.get("respuesta_final")) or "(no provista)"
        des = _coerce_list_str(r.get("desarrollo"), 12)
        ef = _coerce_list_str(r.get("errores_frecuentes"), 8)
        respuestas_out.append({
            "item_index": i,
            "respuesta_final": rf,
            "desarrollo": des,
            "errores_frecuentes": ef,
        })

    crit = _coerce_list_str(sol.get("criterios_correccion"), 20)
    if not crit:
        crit = ["Cumple consigna.", "Legible.", "Completo."]

    data["solucionario_docente"] = {"respuestas": respuestas_out, "criterios_correccion": crit}

    # control_calidad (for display; not used as blocker)
    cc = data.get("control_calidad", {})
    if not isinstance(cc, dict):
        cc = {}
    data["control_calidad"] = {
        "items_count": len(data["items"]),
        "incluye_ejemplo": bool(cc.get("incluye_ejemplo", True)),
        "lenguaje_concreto": bool(cc.get("lenguaje_concreto", True)),
        "una_accion_por_frase": bool(cc.get("una_accion_por_frase", True)),
        "sin_markdown": True,
    }

    # objective/consigna minimums
    if len(data["objetivo_aprendizaje"]) < 12:
        data["objetivo_aprendizaje"] = "Practicar habilidades del tema propuesto."
    if len(data["consigna_general_alumno"]) < 40:
        data["consigna_general_alumno"] = "1. Lee cada consigna.\n2. Haz un paso por vez.\n3. Revisa tu trabajo.\n4. Pide ayuda si no entiendes."

    return data


def _too_thin_activity(data: Dict[str, Any]) -> Tuple[bool, str]:
    try:
        if not isinstance(data, dict):
            return True, "data no es dict"
        obj = _clean_str(data.get("objetivo_aprendizaje"))
        cons = _clean_str(data.get("consigna_general_alumno"))
        items = data.get("items", [])
        if len(obj) < 12:
            return True, "objetivo corto/vacío"
        if len(cons) < 40:
            return True, "consigna corta/vacía"
        if not isinstance(items, list) or len(items) < 6:
            return True, "pocos items"
        nonempty = 0
        for it in items:
            if isinstance(it, dict) and _clean_str(it.get("enunciado")):
                nonempty += 1
        if nonempty < 6:
            return True, "items sin enunciado suficiente"
        return False, "OK"
    except Exception as e:
        return True, f"exception: {e}"


def fallback_activity_strong(alumno: Dict[str, str], input_text: str) -> Dict[str, Any]:
    grupo = alumno.get("grupo", "")
    diag = alumno.get("diagnostico", "")
    tema = (input_text or "").strip()[:160]

    data = {
        "objetivo_aprendizaje": f"Practicar habilidades del tema: {tema}" if tema else "Practicar habilidades básicas del grado.",
        "tiempo_total_min": 60,
        "consigna_general_alumno": "1. Lee cada consigna.\n2. Haz un paso por vez.\n3. Marca lo que completas.\n4. Revisa al final.",
        "items": [
            {
                "tipo": "lectura",
                "enunciado": "📖 Lee 6 palabras y marca las que empiezan igual.",
                "pasos": ["Lee despacio.", "Busca la primera letra.", "Marca 3 palabras."],
                "opciones": ["ma-no", "me-sa", "mi-la", "pa-to", "pe-lo", "pi-no"],
                "respuesta_formato": "multiple_choice",
                "keywords_bold": ["primera letra", "marca"],
                "pista_visual": "Mira solo la primera letra de cada palabra.",
                "visual": {"habilitado": True, "prompt": IMAGE_PROMPT_PREFIX + "letra"},
            },
            {
                "tipo": "completar",
                "enunciado": "✍️ Completa la palabra con la sílaba que falta.",
                "pasos": ["Lee la palabra.", "Dila en voz baja.", "Escribe la sílaba."],
                "opciones": ["ma", "me", "mi", "mo", "mu"],
                "respuesta_formato": "texto_corto",
                "keywords_bold": ["completa", "sílaba"],
                "pista_visual": "Di la palabra y escucha el sonido que falta.",
                "visual": {"habilitado": True, "prompt": IMAGE_PROMPT_PREFIX + "sílaba"},
            },
            {
                "tipo": "escritura",
                "enunciado": "✍️ Escribe 3 oraciones cortas de 3 palabras.",
                "pasos": ["Elige una palabra.", "Escribe una oración.", "Repite 3 veces."],
                "opciones": [],
                "respuesta_formato": "texto_corto",
                "keywords_bold": ["3 oraciones", "cortas"],
                "pista_visual": "Usa: sujeto + verbo + cosa. Ej: “Yo veo sol”.",
                "visual": {"habilitado": True, "prompt": IMAGE_PROMPT_PREFIX + "niño escribiendo"},
            },
            {
                "tipo": "unir",
                "enunciado": "📖 Une palabra con dibujo (mismo significado).",
                "pasos": ["Lee la palabra.", "Busca el dibujo.", "Traza una línea."],
                "opciones": ["sol", "pan", "mesa", "mono"],
                "respuesta_formato": "procedimiento",
                "keywords_bold": ["une", "dibujo"],
                "pista_visual": "Piensa qué objeto es cada palabra.",
                "visual": {"habilitado": True, "prompt": IMAGE_PROMPT_PREFIX + "sol"},
            },
            {
                "tipo": "calcular",
                "enunciado": "🔢 Cuenta objetos y escribe el número.",
                "pasos": ["Cuenta 1 por 1.", "Repite contando otra vez.", "Escribe el número."],
                "opciones": [],
                "respuesta_formato": "texto_corto",
                "keywords_bold": ["cuenta", "número"],
                "pista_visual": "Señala con el dedo cada objeto al contar.",
                "visual": {"habilitado": True, "prompt": IMAGE_PROMPT_PREFIX + "manzanas"},
            },
            {
                "tipo": "dibujar",
                "enunciado": "🎨 Dibuja una escena simple y escribe 2 palabras.",
                "pasos": ["Dibuja 3 cosas.", "Escribe 2 palabras.", "Revisa que se lean."],
                "opciones": [],
                "respuesta_formato": "dibujo",
                "keywords_bold": ["dibuja", "2 palabras"],
                "pista_visual": "Elige cosas simples: sol, casa, árbol.",
                "visual": {"habilitado": True, "prompt": IMAGE_PROMPT_PREFIX + "casa"},
            },
        ],
        "adecuaciones_aplicadas": [
            f"Perfil: {diag}" if diag else "Ajuste por perfil del alumno.",
            f"Grupo: {grupo}" if grupo else "Ajuste por grupo.",
            "Consignas cortas, 1 acción por frase.",
        ],
        "sugerencias_docente": [
            "Modelar 1 ejemplo antes de iniciar.",
            "Pausas breves cada 10–12 min.",
            "Chequeo rápido de comprensión (1 pregunta).",
        ],
        "solucionario_docente": {
            "respuestas": [
                {"item_index": i, "respuesta_final": "(fallback) revisar según ejecución", "desarrollo": ["Ver pasos."], "errores_frecuentes": ["Apuro.", "No sigue pasos."]}
                for i in range(1, 7)
            ],
            "criterios_correccion": ["Cumple consigna.", "Legible.", "Completo."],
        },
        "control_calidad": {"items_count": 6, "incluye_ejemplo": True, "lenguaje_concreto": True, "una_accion_por_frase": True, "sin_markdown": True},
    }
    return normalize_activity_payload(data)


# ============================================================
# Gemini generation (robust)
# ============================================================
def build_prompt(mode: str, grado: str, grupo: str, diagnostico: str, nombre: str, input_text: str) -> str:
    # CLAVE: incluye SIEMPRE grupo + diagnostico + grado por alumno
    if mode == "CREAR":
        ctx = f"CREAR ACTIVIDAD DESDE CERO:\n{input_text}\n"
    else:
        ctx = f"ADAPTAR CONTENIDO ORIGINAL:\n{input_text}\n"

    return f"""{SYSTEM_PROMPT}

CONTEXTO:
{ctx}

ALUMNO (planilla):
- nombre: {nombre}
- diagnostico: {diagnostico}
- grupo: {grupo}
- grado: {grado}

REGLAS EXTRA:
- items[] debe tener EXACTAMENTE 6 items.
- No dejes campos vacíos.
- NO markdown.
- keywords_bold[] corto (2 a 6 palabras).
- visual.habilitado=true y visual.prompt con el prefijo ARASAAC.

DEVUELVE SOLO JSON.
""".strip()


def robust_generate_activity(model_id: str, prompt: str, max_out: int) -> Tuple[Dict[str, Any], str]:
    m = genai.GenerativeModel(model_id)
    cfg = dict(BASE_GEN_CFG_JSON)
    cfg["max_output_tokens"] = max_out

    resp = retry_with_backoff(lambda: m.generate_content(prompt, generation_config=cfg, safety_settings=SAFETY_SETTINGS))
    raw = _extract_text_or_none(resp)
    fr = _finish_reason(resp)
    if raw is None:
        raise ValueError(f"Empty candidate (finish_reason={fr})")

    # 1) parse directo
    try:
        return safe_json_loads(raw), "RAW_PARSE_OK"
    except Exception as pe:
        # 2) parse-fix por modelo (solo sintaxis)
        fix_prompt = build_parse_fix_prompt(raw, f"{type(pe).__name__}: {pe}")
        resp2 = retry_with_backoff(lambda: m.generate_content(fix_prompt, generation_config=cfg, safety_settings=SAFETY_SETTINGS))
        raw2 = _extract_text_or_none(resp2)
        fr2 = _finish_reason(resp2)
        if raw2 is None:
            raise ValueError(f"Empty after parse-fix (finish_reason={fr2})")
        try:
            return safe_json_loads(raw2), "PARSE_FIX_OK"
        except Exception as pe2:
            # 3) repair por reglas
            rep = build_repair_prompt(raw2, f"{type(pe2).__name__}: {pe2}")
            resp3 = retry_with_backoff(lambda: m.generate_content(rep, generation_config=cfg, safety_settings=SAFETY_SETTINGS))
            raw3 = _extract_text_or_none(resp3)
            fr3 = _finish_reason(resp3)
            if raw3 is None:
                raise ValueError(f"Empty after repair (finish_reason={fr3})")
            return safe_json_loads(raw3), "REPAIR_OK"


@st.cache_data(ttl=CACHE_TTL_SECONDS, show_spinner=False)
def cached_activity(cache_key: str, prompt: str, max_out: int) -> Tuple[Dict[str, Any], str]:
    # cache only the raw parse result (still normalized later)
    data, path = robust_generate_activity(DEFAULT_TEXT_MODEL, prompt, max_out=max_out)
    return data, path


def generate_with_retries(cache_key: str, prompt: str) -> Tuple[Dict[str, Any], str, int]:
    last_err = None

    # primary tries (cached)
    for t in OUT_TOKEN_STEPS:
        try:
            data, path = cached_activity(f"{cache_key}::TOK::{t}", prompt, max_out=t)
            return data, f"CACHED::{path}", t
        except Exception as e:
            last_err = e

    # fallback tries (cached)
    for t in OUT_TOKEN_STEPS_FALLBACK:
        try:
            data, path = cached_activity(f"{cache_key}::TOK_FB::{t}", prompt, max_out=t)
            return data, f"CACHED_FB::{path}", t
        except Exception as e:
            last_err = e

    raise last_err if last_err else RuntimeError("Fallo desconocido generando actividad")


# ============================================================
# DOCX rendering
# ============================================================
def apply_card_style(cell, fill_hex: str = "FAFAFA"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)

    tc_borders = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']:
        edge = OxmlElement(f'w:{b}')
        edge.set(qn('w:val'), 'single')
        edge.set(qn('w:sz'), '4')
        edge.set(qn('w:space'), '0')
        edge.set(qn('w:color'), 'E0E0E0')
        tc_borders.append(edge)
    tc_pr.append(tc_borders)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def add_text(paragraph, text: str, bold: bool = False, color: Optional[RGBColor] = None, size_pt: int = 14):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = False
    run.font.name = "Verdana"
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color
    return run


def add_text_with_keywords(paragraph, text: str, keywords: List[str], size_pt: int = 14):
    text = str(text or "")
    kws = [k.strip() for k in (keywords or []) if isinstance(k, str) and k.strip()]
    if not kws:
        add_text(paragraph, text, bold=False, size_pt=size_pt)
        return
    kws_sorted = sorted(set(kws), key=len, reverse=True)
    pat = re.compile("(" + "|".join(re.escape(k) for k in kws_sorted) + ")")
    parts = pat.split(text)
    for part in parts:
        if part in kws_sorted:
            add_text(paragraph, part, bold=True, size_pt=size_pt)
        else:
            add_text(paragraph, part, bold=False, size_pt=size_pt)


def response_box(cell, label: str = "✍️ Respuesta:", lines: int = 4):
    t = cell.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    apply_card_style(c, fill_hex="FFFFFF")
    clear_paragraph(c.paragraphs[0])

    p = c.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    add_text(p, label + " ", bold=True)

    p2 = c.add_paragraph()
    p2.paragraph_format.line_spacing = 1.5
    add_text(p2, "\n" + ("\n" * max(0, lines - 1)) + " ", bold=False)


def checkbox_list(cell, options: List[str], max_opts: int = 8):
    for opt in (options or [])[:max_opts]:
        p = cell.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        add_text(p, f"☐ {opt}", bold=False)


def header_block(doc: Document, alumno: Dict[str, str], logo_b: Optional[bytes], title: str):
    style = doc.styles['Normal']
    style.font.name = 'Verdana'
    style.font.size = Pt(12)

    h = doc.add_table(rows=1, cols=2)
    h.width = Inches(6.5)

    if logo_b:
        try:
            h.rows[0].cells[0].paragraphs[0].add_run().add_picture(io.BytesIO(logo_b), width=Inches(0.7))
        except Exception:
            pass

    info = h.rows[0].cells[1].paragraphs[0]
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(info, title + "\n", bold=True, size_pt=12)
    add_text(info, f"{alumno.get('nombre','')}\n", bold=True, size_pt=11)
    add_text(info, f"{alumno.get('diagnostico','')}\n", bold=False, size_pt=11)
    add_text(info, f"Grupo: {alumno.get('grupo','')} | Grado: {alumno.get('grado','')}", bold=False, size_pt=11)

    doc.add_paragraph("")


def render_alumno_docx(data: Dict[str, Any], alumno: Dict[str, str], logo_b: Optional[bytes], img_model_id: Optional[str], enable_img: bool) -> bytes:
    doc = Document()
    header_block(doc, alumno, logo_b, "FICHA DEL ALUMNO")

    p = doc.add_paragraph()
    add_text(p, "Objetivo de aprendizaje", bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.5
    add_text(p2, str(data.get("objetivo_aprendizaje", "")), bold=False)
    doc.add_paragraph("")

    p = doc.add_paragraph()
    add_text(p, "Consigna general (paso a paso)", bold=True)
    cg = str(data.get("consigna_general_alumno", "")).strip()
    for line in [x.strip() for x in cg.split("\n") if x.strip()]:
        p3 = doc.add_paragraph()
        p3.paragraph_format.line_spacing = 1.4
        add_text(p3, line, bold=False)
    doc.add_paragraph("")

    for idx, it in enumerate(data.get("items", []), start=1):
        if not isinstance(it, dict):
            continue

        tipo = str(it.get("tipo", "")).strip()
        enunciado = ensure_action_emoji(tipo, str(it.get("enunciado", "")).strip())
        pasos = it.get("pasos", []) if isinstance(it.get("pasos", []), list) else []
        opciones = it.get("opciones", []) if isinstance(it.get("opciones", []), list) else []
        formato = str(it.get("respuesta_formato", "texto_corto")).strip()
        kw = it.get("keywords_bold", []) if isinstance(it.get("keywords_bold", []), list) else []
        pista = str(it.get("pista_visual", "")).strip()

        v = it.get("visual", {}) if isinstance(it.get("visual", {}), dict) else {}
        v_en = normalize_bool(v.get("habilitado", False))
        v_pr = normalize_visual_prompt(str(v.get("prompt", "")).strip()) if v_en else ""

        card = doc.add_table(rows=1, cols=1)
        card.width = Inches(6.5)
        cell = card.rows[0].cells[0]
        apply_card_style(cell, fill_hex="FAFAFA")
        clear_paragraph(cell.paragraphs[0])

        pt = cell.add_paragraph()
        pt.paragraph_format.line_spacing = 1.3
        add_text(pt, f"Ítem {idx}", bold=True, size_pt=12)

        p_con = cell.add_paragraph()
        p_con.paragraph_format.line_spacing = 1.5
        add_text_with_keywords(p_con, enunciado, kw, size_pt=13)

        if pasos:
            for i, step in enumerate(pasos[:8], start=1):
                ps = cell.add_paragraph()
                ps.paragraph_format.line_spacing = 1.4
                add_text(ps, f"{i}. {str(step)}", bold=False, size_pt=12)

        sep = cell.add_paragraph()
        add_text(sep, "Trabajo", bold=True, size_pt=12)

        if opciones and (formato.lower() in {"multiple_choice", "multiple choice"}):
            checkbox_list(cell, [str(x) for x in opciones], max_opts=8)
        else:
            response_box(cell, label="✍️ Respuesta:", lines=4)

        if pista:
            pp = cell.add_paragraph()
            pp.paragraph_format.line_spacing = 1.3
            add_text(pp, "Pista", bold=True, size_pt=12)
            pp2 = cell.add_paragraph()
            pp2.paragraph_format.line_spacing = 1.3
            add_text(pp2, "💡 " + pista, bold=False, color=RGBColor(0, 120, 0), size_pt=12)

        if enable_img and img_model_id and v_en and v_pr:
            img_bytes = generate_image_bytes(img_model_id, v_pr)
            if img_bytes:
                pi = cell.add_paragraph()
                add_text(pi, "Apoyo visual", bold=True, size_pt=12)
                pic = cell.add_paragraph()
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    pic.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(2.1))
                except Exception:
                    pass

        doc.add_paragraph("")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def render_docente_docx(data: Dict[str, Any], alumno: Dict[str, str], logo_b: Optional[bytes]) -> bytes:
    doc = Document()
    header_block(doc, alumno, logo_b, "SOLUCIONARIO DOCENTE")

    p = doc.add_paragraph()
    add_text(p, "Objetivo de aprendizaje", bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.4
    add_text(p2, str(data.get("objetivo_aprendizaje", "")), bold=False)
    doc.add_paragraph("")

    sol = data.get("solucionario_docente", {}) if isinstance(data.get("solucionario_docente", {}), dict) else {}
    respuestas = sol.get("respuestas", []) if isinstance(sol.get("respuestas", []), list) else []

    by_idx: Dict[int, Dict[str, Any]] = {}
    for r in respuestas:
        if not isinstance(r, dict):
            continue
        idx = int(r.get("item_index", 0) or 0)
        by_idx[idx] = r

    for idx, it in enumerate(data.get("items", []), start=1):
        if not isinstance(it, dict):
            continue

        en = str(it.get("enunciado", "")).strip()
        r = by_idx.get(idx, {})

        card = doc.add_table(rows=1, cols=1)
        card.width = Inches(6.5)
        cell = card.rows[0].cells[0]
        apply_card_style(cell, fill_hex="FFFFFF")
        clear_paragraph(cell.paragraphs[0])

        pt = cell.add_paragraph()
        pt.paragraph_format.line_spacing = 1.3
        add_text(pt, f"Ítem {idx}", bold=True, size_pt=12)

        pe = cell.add_paragraph()
        pe.paragraph_format.line_spacing = 1.4
        add_text(pe, en, bold=True, size_pt=12)

        pf = cell.add_paragraph()
        pf.paragraph_format.line_spacing = 1.4
        add_text(pf, "Respuesta final: ", bold=True, size_pt=12)
        add_text(pf, str(r.get("respuesta_final", "(no provista)")), bold=False, size_pt=12)

        des = r.get("desarrollo", []) if isinstance(r.get("desarrollo", []), list) else []
        if des:
            pd = cell.add_paragraph()
            add_text(pd, "Desarrollo:", bold=True, size_pt=12)
            for step in des[:12]:
                ps = cell.add_paragraph()
                ps.paragraph_format.line_spacing = 1.3
                add_text(ps, f"• {step}", bold=False, size_pt=12)

        ef = r.get("errores_frecuentes", []) if isinstance(r.get("errores_frecuentes", []), list) else []
        if ef:
            pef = cell.add_paragraph()
            add_text(pef, "Errores frecuentes:", bold=True, size_pt=12)
            for e in ef[:8]:
                pex = cell.add_paragraph()
                pex.paragraph_format.line_spacing = 1.3
                add_text(pex, f"• {e}", bold=False, size_pt=12)

        doc.add_paragraph("")

    crit = sol.get("criterios_correccion", []) if isinstance(sol.get("criterios_correccion", []), list) else []
    if crit:
        p = doc.add_paragraph()
        add_text(p, "Criterios de corrección", bold=True, size_pt=12)
        for c in crit[:15]:
            p2 = doc.add_paragraph()
            p2.paragraph_format.line_spacing = 1.3
            add_text(p2, f"• {c}", bold=False, size_pt=12)

    adec = data.get("adecuaciones_aplicadas", []) if isinstance(data.get("adecuaciones_aplicadas", []), list) else []
    if adec:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        add_text(p, "Adecuaciones aplicadas", bold=True, size_pt=12)
        for a in adec[:20]:
            p2 = doc.add_paragraph()
            p2.paragraph_format.line_spacing = 1.3
            add_text(p2, f"• {a}", bold=False, size_pt=12)

    sug = data.get("sugerencias_docente", []) if isinstance(data.get("sugerencias_docente", []), list) else []
    if sug:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        add_text(p, "Sugerencias para el docente", bold=True, size_pt=12)
        for s in sug[:20]:
            p2 = doc.add_paragraph()
            p2.paragraph_format.line_spacing = 1.3
            add_text(p2, f"• {s}", bold=False, size_pt=12)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ============================================================
# UI + Process
# ============================================================
def main():
    st.title("Nano Opal v25.0 🍌")
    st.caption("Modelo fijo gemini-2.5-flash. Sin selectors. JSON tolerante + fallback anti-archivos vacíos.")

    # API key
    try:
        genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])
    except Exception as e:
        st.error(f"Falta/invalid GOOGLE_API_KEY en secrets: {e}")
        return

    # Smoke test text
    ok_sm, msg_sm = smoke_test_text_model(DEFAULT_TEXT_MODEL)
    if not ok_sm:
        st.error(f"Texto no responde ({DEFAULT_TEXT_MODEL}): {msg_sm}")
        return

    # Load sheet
    try:
        df = pd.read_csv(URL_PLANILLA)
        df.columns = [c.strip() for c in df.columns]
    except Exception as e:
        st.error(f"Error cargando planilla: {e}")
        return

    # Column mapping: mantengo tu heurística
    grado_col = df.columns[1] if len(df.columns) > 1 else df.columns[0]
    alumno_col = df.columns[2] if len(df.columns) > 2 else df.columns[0]
    grupo_col = df.columns[3] if len(df.columns) > 3 else df.columns[0]
    diag_col = df.columns[4] if len(df.columns) > 4 else df.columns[0]

    # Sidebar minimal (sin selector de modelos / sin on-the-fly)
    with st.sidebar:
        st.header("📚 Lote")
        grado = st.selectbox("Grado", sorted(df[grado_col].dropna().unique().tolist()))
        df_g = df[df[grado_col] == grado].copy()

        alcance = st.radio("Alcance", ["Todo el grado", "Seleccionar alumnos"], horizontal=True)
        if alcance == "Seleccionar alumnos":
            al_sel = st.multiselect("Alumnos", sorted(df_g[alumno_col].dropna().unique().tolist()))
            df_final = df_g[df_g[alumno_col].isin(al_sel)].copy() if al_sel else df_g.iloc[0:0].copy()
        else:
            df_final = df_g

        st.divider()
        enable_img = st.checkbox("Habilitar imágenes", value=True)
        logo = st.file_uploader("Logo", type=["png", "jpg", "jpeg"])
        l_bytes = logo.read() if logo else None

        st.divider()
        if st.button("Limpiar cache"):
            st.cache_data.clear()
            st.cache_resource.clear()
            st.success("Cache limpia.")

    # Tabs: Adaptar vs Crear
    tab1, tab2 = st.tabs(["🔄 Adaptar DOCX", "✨ Crear Actividad"])

    adapt_docx = None
    if "brief_committed" not in st.session_state:
        st.session_state["brief_committed"] = ""

    with tab1:
        st.subheader("Adaptar (DOCX)")
        adapt_docx = st.file_uploader("Actividad base (DOCX)", type=["docx"], key="docx_in")

    with tab2:
        st.subheader("Crear desde prompt (con botón)")
        with st.form("form_prompt", clear_on_submit=False):
            brief_draft = st.text_area(
                "Prompt",
                height=220,
                key="brief_draft",
                placeholder="Tema + objetivo + grado. Ej: 1ero: sílabas directas + lectura de palabras simples."
            )
            submitted = st.form_submit_button("✅ Enviar prompt")
        if submitted:
            st.session_state["brief_committed"] = (brief_draft or "").strip()
            st.success("Prompt enviado.")

        if st.session_state["brief_committed"]:
            st.info("Listo: hay un prompt enviado para el lote.")

    # Determine mode + input_text
    brief = (st.session_state.get("brief_committed") or "").strip()
    mode = "CREAR" if brief else "ADAPTAR"

    input_text = ""
    if mode == "ADAPTAR":
        if adapt_docx:
            input_text = extraer_texto_docx(adapt_docx)
            ok_in, msg_in, info_in = validate_text_input(input_text, "ADAPTAR")
            if ok_in:
                st.success(f"Parseo DOCX OK ({info_in['chars']} chars)")
            else:
                st.error(f"Parseo DOCX: {msg_in}")
            with st.expander("Preview texto extraído", expanded=False):
                st.text(info_in.get("preview", ""))
        else:
            st.warning("Subí un DOCX o ve a 'Crear Actividad' y enviá un prompt.")
            return
    else:
        input_text = brief
        ok_in, msg_in, info_in = validate_text_input(input_text, "CREAR")
        if ok_in:
            st.success(f"Prompt OK ({info_in['chars']} chars)")
        else:
            st.error(f"Prompt: {msg_in}")
            return
        with st.expander("Preview prompt enviado", expanded=False):
            st.text(info_in.get("preview", ""))

    # Resolve image model auto
    img_model_id = pick_image_model() if enable_img else None
    if enable_img and not img_model_id:
        st.warning("Imágenes activadas, pero no se detectó modelo de imagen compatible. Se omiten imágenes.")

    # Generate button
    if st.button("🚀 GENERAR LOTE"):
        if len(df_final) == 0:
            st.error("No hay alumnos (ver selección por grado/alumnos).")
            return

        ok_in, msg_in, _ = validate_text_input(input_text, mode)
        if not ok_in:
            st.error(f"No se inicia: {msg_in}")
            return

        zip_io = io.BytesIO()
        ok_count = 0
        err_count = 0
        errors: List[str] = []

        logs = []
        logs.append("Nano Opal v25.0")
        logs.append(f"Inicio: {now_str()}")
        logs.append(f"Modo: {mode}")
        logs.append(f"Grado: {grado}")
        logs.append(f"Modelo texto: {DEFAULT_TEXT_MODEL}")
        logs.append(f"Modelo imagen: {img_model_id if img_model_id else 'N/A'}")
        logs.append(f"Imagen habilitada: {bool(enable_img and img_model_id)}")
        logs.append(f"Alumnos: {len(df_final)}")
        logs.append("")

        resumen = []
        resumen.append(f"RESUMEN - Nano Opal v25.0")
        resumen.append(f"Inicio: {now_str()}")
        resumen.append(f"Modo: {mode}")
        resumen.append(f"Grado: {grado}")
        resumen.append("")

        prog = st.progress(0.0)
        status = st.empty()

        base_hash = hash_text(f"{mode}|{grado}|{input_text}|{DEFAULT_TEXT_MODEL}|v25")

        with zipfile.ZipFile(zip_io, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("_REPORTE.txt", "\n".join(logs))

            for irow, (_, row) in enumerate(df_final.iterrows(), start=1):
                nombre = str(row[alumno_col]).strip()
                grupo = str(row[grupo_col]).strip()
                diagnostico = str(row[diag_col]).strip()

                alumno_meta = {
                    "nombre": nombre,
                    "grupo": grupo,
                    "grado": str(grado),
                    "diagnostico": diagnostico,
                }

                status.info(f"Procesando: {nombre} ({irow}/{len(df_final)})")
                prog.progress(min(1.0, irow / max(1, len(df_final))))

                try:
                    prompt = build_prompt(mode, str(grado), grupo, diagnostico, nombre, input_text)
                    cache_key = f"{base_hash}::{safe_filename(nombre)}::{safe_filename(grupo)}::{safe_filename(diagnostico)}"

                    # 1) generación principal (cacheada)
                    data_raw, gen_path, tok = generate_with_retries(cache_key, prompt)
                    data = normalize_activity_payload(data_raw)

                    # 2) anti-thin -> reintento sin cache (cache-bust)
                    thin, reason = _too_thin_activity(data)
                    if thin:
                        prompt2 = prompt + "\n\nOBLIGATORIO: 6 items completos. Campos NO vacíos. consigna >= 4 líneas."
                        data_raw2, gen_path2 = robust_generate_activity(DEFAULT_TEXT_MODEL, prompt2, max_out=4096)
                        data2 = normalize_activity_payload(data_raw2)
                        thin2, reason2 = _too_thin_activity(data2)

                        if not thin2:
                            data = data2
                            gen_path = f"{gen_path} -> RETRY({gen_path2})"
                        else:
                            data = fallback_activity_strong(alumno_meta, input_text)
                            gen_path = f"{gen_path} -> RETRY_THIN({reason2}) -> FALLBACK_STRONG"

                    # 3) render
                    alumno_docx_b = render_alumno_docx(
                        data=data,
                        alumno=alumno_meta,
                        logo_b=l_bytes,
                        img_model_id=img_model_id,
                        enable_img=bool(enable_img and img_model_id),
                    )
                    docente_docx_b = render_docente_docx(
                        data=data,
                        alumno=alumno_meta,
                        logo_b=l_bytes,
                    )

                    base = f"{safe_filename(nombre)}__{safe_filename(grupo)}__{safe_filename(grado)}"
                    zf.writestr(f"{base}__ALUMNO.docx", alumno_docx_b)
                    zf.writestr(f"{base}__DOCENTE.docx", docente_docx_b)

                    ok_count += 1
                    resumen.append(f"- {nombre} | OK ({gen_path})")

                except Exception as e:
                    err_count += 1
                    msg = f"{nombre} | ERROR {type(e).__name__}: {e}"
                    errors.append(msg)
                    resumen.append(f"- {nombre} | ERROR {type(e).__name__}: {e}")

            # resumen final
            resumen.append("")
            resumen.append(f"OK: {ok_count}/{len(df_final)}")
            resumen.append(f"Errores: {err_count}")

            if errors:
                resumen.append("")
                resumen.append("DETALLE ERRORES:")
                resumen.extend(errors[:200])

            zf.writestr("_RESUMEN.txt", "\n".join(resumen))

        status.success(f"Listo. OK={ok_count} / {len(df_final)} | Errores={err_count}")

        zip_io.seek(0)
        st.download_button(
            "⬇️ Descargar ZIP",
            data=zip_io.getvalue(),
            file_name=f"Nano_Opal_v25_{safe_filename(grado)}_{datetime.now().strftime('%Y%m%d_%H%M')}.zip",
            mime="application/zip",
        )


if __name__ == "__main__":
    main()
